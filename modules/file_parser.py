import os
from docx import Document
import fitz  # PyMuPDF

class DocxParser:
    """Word文档解析类"""
    
    def __init__(self, file_path):
        """
        初始化Word文档解析器
        
        参数:
            file_path (str): Word文档的路径
        """
        self.file_path = file_path
        
    def extract_content(self, chapter=None):
        """
        提取Word文档内容
        
        参数:
            chapter (str, optional): 指定章节名称，如果提供则只提取该章节内容
            
        返回:
            str: 提取的文本内容
        """
        try:
            doc = Document(self.file_path)
            
            # 如果指定了章节，尝试提取该章节
            if chapter:
                content = []
                found_chapter = False
                
                for para in doc.paragraphs:
                    # 检查是否是章节标题
                    if chapter.lower() in para.text.lower() and para.style.name.startswith('Heading'):
                        found_chapter = True
                        content.append(para.text)
                    # 如果已找到章节，继续添加内容直到下一个标题
                    elif found_chapter:
                        if para.style.name.startswith('Heading'):
                            break
                        content.append(para.text)
                
                return '\n'.join(content)
            
            # 如果没有指定章节，提取全文
            else:
                return '\n'.join([para.text for para in doc.paragraphs])
                
        except Exception as e:
            return f"解析Word文档时出错: {str(e)}"
    
    def extract_all_chapters(self):
        """
        提取Word文档中的所有章节
        
        返回:
            dict: 章节名称和内容的字典
        """
        try:
            doc = Document(self.file_path)
            chapters = {}
            current_chapter = "前言"
            chapter_content = []
            
            for para in doc.paragraphs:
                # 检查是否是章节标题
                if para.style.name.startswith('Heading'):
                    # 保存上一章节内容
                    if chapter_content:
                        chapters[current_chapter] = '\n'.join(chapter_content)
                    
                    # 开始新章节
                    current_chapter = para.text
                    chapter_content = [para.text]
                else:
                    chapter_content.append(para.text)
            
            # 保存最后一章节内容
            if chapter_content:
                chapters[current_chapter] = '\n'.join(chapter_content)
                
            return chapters
            
        except Exception as e:
            return {"错误": f"提取章节时出错: {str(e)}"}


class PdfParser:
    """PDF文档解析类"""
    
    def __init__(self, file_path):
        """
        初始化PDF文档解析器
        
        参数:
            file_path (str): PDF文档的路径
        """
        self.file_path = file_path
        
    def extract_content(self, page_range=None):
        """
        提取PDF文档内容
        
        参数:
            page_range (tuple, optional): 页面范围，如(0, 5)表示提取前5页
            
        返回:
            str: 提取的文本内容
        """
        try:
            doc = fitz.open(self.file_path)
            
            # 如果指定了页面范围
            if page_range:
                start, end = page_range
                end = min(end, len(doc))
                content = []
                
                for page_num in range(start, end):
                    page = doc[page_num]
                    content.append(page.get_text())
                    
                return '\n'.join(content)
            
            # 如果没有指定页面范围，提取全文
            else:
                content = []
                for page in doc:
                    content.append(page.get_text())
                
                return '\n'.join(content)
                
        except Exception as e:
            return f"解析PDF文档时出错: {str(e)}"
    
    def extract_by_keywords(self, keywords):
        """
        根据关键词提取PDF文档中的相关内容
        
        参数:
            keywords (list): 关键词列表
            
        返回:
            dict: 关键词和相关内容的字典
        """
        try:
            doc = fitz.open(self.file_path)
            results = {}
            
            for keyword in keywords:
                keyword_results = []
                
                for page_num, page in enumerate(doc):
                    # 搜索关键词
                    text_instances = page.search_for(keyword)
                    
                    if text_instances:
                        # 获取页面文本
                        page_text = page.get_text()
                        keyword_results.append(f"页面 {page_num+1}: {page_text}")
                
                if keyword_results:
                    results[keyword] = '\n'.join(keyword_results)
                else:
                    results[keyword] = "未找到相关内容"
            
            return results
            
        except Exception as e:
            return {"错误": f"按关键词提取内容时出错: {str(e)}"}
