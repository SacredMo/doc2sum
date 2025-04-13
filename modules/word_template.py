from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

class WordTemplate:
    """Word模板填充类"""
    
    def __init__(self, template_path=None):
        """
        初始化Word模板填充器
        
        参数:
            template_path (str, optional): Word模板文件的路径，如果不提供则创建默认模板
        """
        self.template_path = template_path
        self.output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
        
        # 确保输出目录存在
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # 如果没有提供模板，创建默认模板
        if not template_path or not os.path.exists(template_path):
            self.create_default_template()
    
    def create_default_template(self):
        """创建默认的Word模板"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading("文件内容总结报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加说明段落
        doc.add_paragraph("本报告由文件处理平台自动生成，包含文件内容的AI辅助总结。")
        doc.add_paragraph()
        
        # 创建表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "日期时间"
        header_cells[1].text = "文件名称"
        header_cells[2].text = "指定内容"
        header_cells[3].text = "AI总结内容"
        
        # 保存模板
        self.template_path = os.path.join(self.output_dir, "default_template.docx")
        doc.save(self.template_path)
    
    def fill_template(self, file_name, specified_content, summary_content):
        """
        填充Word模板
        
        参数:
            file_name (str): 文件名称
            specified_content (str): 指定的章节或关键内容
            summary_content (str): AI总结内容
            
        返回:
            str: 生成的Word文档路径
        """
        try:
            # 加载模板
            doc = Document(self.template_path)
            
            # 查找表格
            if len(doc.tables) > 0:
                table = doc.tables[0]
                
                # 添加新行
                row = table.add_row()
                cells = row.cells
                
                # 填充单元格内容
                cells[0].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                cells[1].text = file_name
                cells[2].text = specified_content if specified_content else "无指定内容"
                cells[3].text = summary_content
            else:
                # 如果没有找到表格，添加内容到文档末尾
                doc.add_heading(f"文件: {file_name}", level=1)
                
                if specified_content:
                    doc.add_heading("指定内容", level=2)
                    doc.add_paragraph(specified_content)
                
                doc.add_heading("AI总结内容", level=2)
                doc.add_paragraph(summary_content)
                
                doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_file = os.path.join(self.output_dir, f"summary_{timestamp}.docx")
            doc.save(output_file)
            
            return output_file
            
        except Exception as e:
            print(f"填充Word模板时出错: {str(e)}")
            return None
    
    def create_report(self, records, output_file=None):
        """
        根据多条记录创建综合报告
        
        参数:
            records (list): 记录列表，每条记录应包含file_name、specified_content和summary字段
            output_file (str, optional): 输出文件路径，如果不提供则自动生成
            
        返回:
            str: 生成的Word文档路径
        """
        try:
            # 创建新文档
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading("文件内容总结综合报告", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加说明段落
            doc.add_paragraph("本报告由文件处理平台自动生成，包含多个文件内容的AI辅助总结。")
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # 添加每条记录
            for i, record in enumerate(records):
                doc.add_heading(f"文件 {i+1}: {record['file_name']}", level=1)
                
                if record.get('specified_content'):
                    doc.add_heading("指定内容", level=2)
                    doc.add_paragraph(record['specified_content'])
                
                doc.add_heading("AI总结内容", level=2)
                doc.add_paragraph(record['summary'])
                
                # 添加分隔线（除了最后一条记录）
                if i < len(records) - 1:
                    doc.add_paragraph("---")
            
            # 确定输出文件路径
            if not output_file:
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                output_file = os.path.join(self.output_dir, f"comprehensive_report_{timestamp}.docx")
            
            # 保存文档
            doc.save(output_file)
            
            return output_file
            
        except Exception as e:
            print(f"创建综合报告时出错: {str(e)}")
            return None
